import requests
from bs4 import BeautifulSoup
import time
import random
import fitz  # PyMuPDF
import docx  # python-docx
from scholarly import scholarly
import re  # Para buscar palabras clave
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from tenacity import retry, stop_after_attempt, wait_fixed
from azure.ai.inference import ChatCompletionsClient
from azure.core.credentials import AzureKeyCredential
from azure.ai.inference.models import SystemMessage, UserMessage
import os
from docx import Document

# Configuración del cliente de Azure
model = ChatCompletionsClient(
    endpoint="END_POINT",  # Cambiar por el endpoint de Azure
    credential=AzureKeyCredential("API_KEY"),
)

# Configurar un proxy personalizado
# Configurar proxies mediante variables de entorno
#os.environ['HTTP_PROXY'] = 'http://your-proxy-ip:port'
#os.environ['HTTPS_PROXY'] = 'https://your-proxy-ip:port'


# Función para resumir texto con sumy
def summarize_text_with_sumy(text, sentences_count=2):
    parser = PlaintextParser.from_string(text, Tokenizer("spanish"))  # Cambia "spanish" al idioma que necesites
    summarizer = LsaSummarizer()
    summary = summarizer(parser.document, sentences_count)
    return " ".join(str(sentence) for sentence in summary)

# Decorador para reintentar en caso de error
@retry(stop=stop_after_attempt(5), wait=wait_fixed(2))  # Reintentar hasta 5 veces, con 2 segundos de espera
def google_scholar_search(query):
    search_query = scholarly.search_pubs(query)
    links = []
    for result in search_query:
        if 'eprint' in result:
            links.append(result['eprint'])
        elif 'pub_url' in result:
            links.append(result['pub_url'])
    return links[:5]  # Limitar a los primeros 5 enlaces

# Función para extraer texto de un PDF
def extract_text_from_pdf(url):
    response = requests.get(url)
    with open("temp.pdf", "wb") as f:
        f.write(response.content)
    doc = fitz.open("temp.pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# Función para extraer texto de un Word
def extract_text_from_word(url):
    response = requests.get(url)
    with open("temp.docx", "wb") as f:
        f.write(response.content)
    doc = docx.Document("temp.docx")
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

# Función para extraer texto de una página web
def extract_text_from_web(url):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36'
    }
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')
    paragraphs = soup.find_all('p')
    text = " ".join([para.get_text() for para in paragraphs])
    return text

# Función para categorizar el texto
def categorize_text(text, categories):
    categorized_text = {category: [] for category in categories}
    for category in categories:
        # Buscar palabras clave relacionadas con la categoría
        keywords = {
            "Aplicaciones prácticas": ["aplicación", "uso", "implementación", "caso de uso"],
            "Tecnologías utilizadas": ["tecnología", "herramienta", "framework", "API"],
            "Desafíos y limitaciones": ["desafío", "limitación", "problema", "retro"],
            "Tendencias futuras": ["futuro", "tendencia", "evolución", "perspectiva"]
        }
        for keyword in keywords[category]:
            if re.search(rf"\b{keyword}\b", text, re.IGNORECASE):
                categorized_text[category].append(text)
                break
    return categorized_text

# Función para resumir el texto
def summarize_text(text, max_length=150):
    try:
        return summarize_text_with_sumy(text, sentences_count=2)  # Resumir en 2 oraciones
    except Exception as e:
        print(f"Error summarizing text: {e}")
        return text[:max_length]  # Si falla, devolver un fragmento del texto

# Función para preparar el prompt
def prepare_prompt(categorized_text):
    prompt = "Información relevante sobre el tema:\n"
    for category, texts in categorized_text.items():
        prompt += f"\n**{category}**:\n"
        for text in texts:
            summary = summarize_text(text)
            prompt += f"- {summary}\n"
    return prompt

# Función principal para extraer texto de los enlaces
def scrape_text_from_links(links):
    all_texts = []
    for link in links:
        print(f"Extracting content from: {link}")
        try:
            if link.endswith('.pdf'):
                text = extract_text_from_pdf(link)
            elif link.endswith('.docx'):
                text = extract_text_from_word(link)
            else:
                text = extract_text_from_web(link)
            all_texts.append(text)
            time.sleep(random.randint(2, 5))  # Simular un pequeño retraso
        except Exception as e:
            print(f"Error extracting text from {link}: {e}")
    return all_texts

# Decorador para reintentar en caso de error 429
@retry(stop=stop_after_attempt(5), wait=wait_fixed(2))  # Reintentar hasta 5 veces, con 2 segundos de espera
def send_request_to_azure(prompt):
    response = model.complete(
        messages=[
            SystemMessage(content="You are a TI Analyst and Doctor in Software Engineer."),
            UserMessage(content=prompt),
        ],
        model_extras={
            "safe_mode": True
        }
    )
    return response

# Categorías predefinidas
categories = [
    "Aplicaciones prácticas",
    "Tecnologías utilizadas",
    "Desafíos y limitaciones",
    "Tendencias futuras"
]

# Búsqueda en Google Scholar
query = "Apis o manejo de rutas publicas usando sistemas de geolocalización o asistentes virtuales para movilizarse en la ciudad"
try:
    links = google_scholar_search(query)
except Exception as e:
    print(f"Error fetching from Google Scholar: {e}")
    links = []  # Continuar con una lista vacía si falla

# Extraer texto de los enlaces
texts = scrape_text_from_links(links)

# Categorizar y resumir los textos
categorized_texts = {category: [] for category in categories}
for text in texts:
    categorized = categorize_text(text, categories)
    for category, fragments in categorized.items():
        categorized_texts[category].extend(fragments)

# Preparar el prompt
prompt = prepare_prompt(categorized_texts)

# Guardar los textos en un archivo TXT
with open("extracted_texts.txt", "w", encoding="utf-8") as f:
    for idx, text in enumerate(texts):
        f.write(f"Text {idx + 1}:\n{text}\n\n")

# Mostrar los textos extraídos
for idx, text in enumerate(texts):
    print(f"\nText {idx + 1}:")
    print(text[:500])  # Mostrar los primeros 500 caracteres como ejemplo

# Crear el prompt para el modelo de Azure
alert = r"Apis o manejo de rutas publicas usando sistemas de geolocalización o asistentes virtuales para movilizarse en la ciudad"
full_prompt = f"""
Actuando como un experto en TI e investigador, del siguiente tema:
{alert}.
Realiza un estado del arte con todos los parámetros que debería llevar un paper de investigación.
La información relevante es:
{prompt}
"""

try:
    response = send_request_to_azure(full_prompt)
    print(response.choices[0].message.content)

    # Guardar la respuesta en un archivo Word
    if not os.path.exists("informe"):
        os.makedirs("informe")

    file_path = os.path.join("informe", "informe_equipo_dinamita.docx")
    doc = Document()
    doc.add_paragraph(response.choices[0].message.content)
    doc.save(file_path)
    print(f"El informe se ha guardado en: {file_path}")

except Exception as e:
    print(f"Error after multiple retries: {e}")